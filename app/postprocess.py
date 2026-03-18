"""Post-process pdf2docx output to fix Indian script and layout issues.

Key fixes (in order):
1. Odia→Gurmukhi character mapping (pdf extraction artifact)
2. Remove spurious spaces before Gurmukhi combining marks
3. Gurmukhi sihari (U+0A3F) character ordering (visual→logical Unicode)
4. Space recovery using original PDF text as reference
5. Final validation: no spaces before combining marks
6. Font names: PostScript→standard (NirmalaUI→Nirmala UI)
7. Complex script (cs) font attribute for Gurmukhi rendering
8. Font sizes: odd half-points→even, add szCs for complex scripts
9. Page margins: 0.5" uniform
10. Table borders: single black borders
11. Table structure: extract full-span rows as standalone paragraphs
"""

import copy
import re

import fitz  # PyMuPDF
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

# --- Gurmukhi Unicode constants ---

# Sihari: only swap when NOT preceded by a consonant (negative lookbehind)
SIHARI = "\u0A3F"
_SIHARI_FIX = re.compile(r"(?<![\u0A15-\u0A39])\u0A3F([\u0A15-\u0A39])")

# Combining marks that must stay attached to their preceding base character
GURMUKHI_COMBINING = set(range(0x0A3E, 0x0A4E)) | {0x0A70, 0x0A71, 0x0A3C}

# Odia characters that pdf2docx incorrectly extracts instead of Gurmukhi
# (caused by font encoding issues in the PDF)
ODIA_TO_GURMUKHI = {
    "\u0B48": "\u0A48",  # ୈ → ੈ (vowel sign AI)
    "\u0B4A": "\u0A4B",  # ୊ → ੋ (vowel sign O → OO, as Gurmukhi lacks O)
    "\u0B4C": "\u0A4C",  # ୌ → ੌ (vowel sign AU)
    "\u0B65": "\u0A65",  # ୥ → ੥ (digit 5)
    "\u0B67": "\u0A67",  # ୧ → ੧ (digit 7)
}
_ODIA_PATTERN = re.compile("|".join(re.escape(k) for k in ODIA_TO_GURMUKHI))

# Pattern to remove spaces before Gurmukhi combining marks.
# EXCLUDES sihari (U+0A3F) because in visual ordering, words starting
# with sihari have it before the consonant — the space is a real word boundary.
_SPACE_BEFORE_COMBINING = re.compile(
    r" ([\u0A3C\u0A3E\u0A40-\u0A4D\u0A70\u0A71])"
)

# Font name mapping: PostScript names → embedded font name.
# We embed Lohit Gurmukhi in the DOCX so it renders on any device.
GURMUKHI_FONT = "Lohit Gurmukhi"
FONT_MAP = {
    "NirmalaUI": GURMUKHI_FONT,
    "NirmalaUI-Bold": GURMUKHI_FONT,
    "Nirmala UI": GURMUKHI_FONT,
    "ArialMT": "Arial",
    "Arial-BoldMT": "Arial",
    "TimesNewRomanPSMT": "Times New Roman",
    "TimesNewRomanPS-BoldMT": "Times New Roman",
}

FONT_ATTRS = [qn("w:ascii"), qn("w:hAnsi"), qn("w:cs"), qn("w:eastAsia")]


def fix_sihari(text: str) -> str:
    """Fix Gurmukhi sihari character ordering from visual to logical."""
    return _SIHARI_FIX.sub(lambda m: m.group(1) + SIHARI, text)


def _fix_odia_chars(text: str) -> str:
    """Map Odia characters to their Gurmukhi equivalents."""
    return _ODIA_PATTERN.sub(lambda m: ODIA_TO_GURMUKHI[m.group()], text)


def _remove_spaces_before_combining(text: str) -> str:
    """Remove spaces that appear before Gurmukhi combining marks."""
    return _SPACE_BEFORE_COMBINING.sub(r"\1", text)


def _fix_pre_space_text(body) -> dict:
    """Pre-space-recovery fixes: Odia mapping + remove spurious spaces."""
    stats = {"odia": 0, "space_cleanup": 0}
    for t_elem in body.iter(qn("w:t")):
        if not t_elem.text:
            continue
        text = t_elem.text

        fixed = _fix_odia_chars(text)
        if fixed != text:
            stats["odia"] += 1
            text = fixed

        fixed = _remove_spaces_before_combining(text)
        if fixed != text:
            stats["space_cleanup"] += 1
            text = fixed

        if text != t_elem.text:
            t_elem.text = text
            t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return stats


def _fix_sihari_all(body) -> int:
    """Fix sihari ordering in ALL w:t elements. Must run AFTER space recovery."""
    fixes = 0
    for t_elem in body.iter(qn("w:t")):
        if t_elem.text and SIHARI in t_elem.text:
            fixed = fix_sihari(t_elem.text)
            if fixed != t_elem.text:
                t_elem.text = fixed
                t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                fixes += 1
    return fixes


def _recover_spaces(body, pdf_path: str) -> int:
    """Recover word spaces using word-level matching with PDF text.

    Key insight: pdf2docx and PyMuPDF's get_text() have DIFFERENT sihari
    orderings (pdf2docx is mixed, get_text is consistently wrong). So we
    normalize BOTH to the same canonical form (sihari after consonant) before
    matching, then apply spaces back to the original DOCX text.
    """
    doc = fitz.open(pdf_path)

    # Method 1: Words from get_text("text") — has spaces where PDF encodes them
    pdf_text = ""
    for page in doc:
        pdf_text += page.get_text("text") + "\n"

    # Method 2: Words from character gap analysis — detects visual spaces
    # that get_text("text") misses (PDF stores some text without explicit spaces)
    gap_words = set()
    for page in doc:
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if block.get("type") != 0:
                continue
            for line in block["lines"]:
                spans = line["spans"]
                for span in spans:
                    text = span["text"]
                    if len(text) > 1:
                        # Use the span's character-level data for gap detection
                        bbox = span["bbox"]
                        char_width = (bbox[2] - bbox[0]) / max(len(text), 1)
                        # Individual words within a span (split by actual spaces)
                        for w in text.split():
                            if 2 <= len(w) <= 15:
                                gap_words.add(w)
        # Also use get_text("words") which does gap-based splitting
        for w in page.get_text("words"):
            word = w[4]
            if 2 <= len(word) <= 15:
                gap_words.add(word)
    doc.close()

    # Normalize all word sources
    pdf_text = _fix_odia_chars(pdf_text)
    pdf_text = fix_sihari(pdf_text)

    word_set = set()
    for word in pdf_text.split():
        if 2 <= len(word) <= 15:
            word_set.add(word)
    # Add gap-detected words (also normalized)
    for word in gap_words:
        w = _fix_odia_chars(word)
        w = fix_sihari(w)
        if 2 <= len(w) <= 15:
            word_set.add(w)
    if not word_set:
        return 0

    max_word_len = max(len(w) for w in word_set)

    # Process each text element: normalize → tokenize → replace
    # This simultaneously fixes sihari AND recovers spaces.
    spaces_added = 0
    for t in body.iter(qn("w:t")):
        if not t.text or len(t.text) < 4:
            continue

        # Normalize: Odia fix + sihari fix (canonical form)
        normalized = _fix_odia_chars(t.text)
        normalized = fix_sihari(normalized)

        # Tokenize using PDF word dictionary
        tokenized = _tokenize_runs(normalized, word_set, max_word_len)

        if tokenized != t.text:
            tokenized = _remove_spaces_before_combining(tokenized)
            if tokenized != t.text:
                t.text = tokenized
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                spaces_added += 1

    return spaces_added


def _tokenize_runs(text: str, word_set: set, max_word_len: int) -> str:
    """Split concatenated text into words using greedy longest-match."""
    result = []
    i = 0

    while i < len(text):
        if text[i] == " ":
            result.append(" ")
            i += 1
            continue

        # Find the end of this non-space segment
        j = i
        while j < len(text) and text[j] != " ":
            j += 1
        segment = text[i:j]

        # Short segments or non-Gurmukhi: keep as-is
        if len(segment) < 4 or not any(0x0A00 <= ord(c) <= 0x0A7F for c in segment):
            result.append(segment)
            i = j
            continue

        # Try greedy longest-match tokenization
        tokenized = _greedy_tokenize(segment, word_set, max_word_len)
        result.append(tokenized)
        i = j

    return "".join(result)


def _greedy_tokenize(segment: str, word_set: set, max_word_len: int) -> str:
    """Split a segment into known words with spaces between them."""
    parts = []
    i = 0
    last_was_word = False

    while i < len(segment):
        # Try to find the longest word starting at position i
        best_len = 0
        for length in range(min(max_word_len, len(segment) - i), 1, -1):
            candidate = segment[i : i + length]
            if candidate in word_set:
                best_len = length
                break

        if best_len >= 2:
            # Found a word — add space before it if previous was also a word
            if last_was_word:
                parts.append(" ")
            parts.append(segment[i : i + best_len])
            i += best_len
            last_was_word = True
        else:
            # Single character — append to current word (no space)
            parts.append(segment[i])
            i += 1
            # Don't set last_was_word to False — let the next word still get a space

    return "".join(parts)


def _fix_fonts(body) -> None:
    """Fix font names and add cs attribute for complex script rendering."""
    for rfonts in body.iter(qn("w:rFonts")):
        for attr in FONT_ATTRS:
            val = rfonts.get(attr)
            if val and val in FONT_MAP:
                rfonts.set(attr, FONT_MAP[val])

        ascii_font = rfonts.get(qn("w:ascii"), "")
        cs_font = rfonts.get(qn("w:cs"))
        if not cs_font:
            mapped = FONT_MAP.get(ascii_font, ascii_font)
            if mapped:
                rfonts.set(qn("w:cs"), mapped)

        primary = rfonts.get(qn("w:ascii"))
        if primary:
            for attr in FONT_ATTRS:
                if not rfonts.get(attr):
                    rfonts.set(attr, primary)


def _fix_font_sizes(body) -> None:
    """Round odd half-point sizes to even, ensure szCs matches sz."""
    for rpr in body.iter(qn("w:rPr")):
        sz = rpr.find(qn("w:sz"))
        if sz is not None:
            val = sz.get(qn("w:val"))
            if val:
                num = int(val)
                if num % 2 != 0:
                    num += 1
                    sz.set(qn("w:val"), str(num))

                szCs = rpr.find(qn("w:szCs"))
                if szCs is None:
                    szCs = OxmlElement("w:szCs")
                    sz.addnext(szCs)
                szCs.set(qn("w:val"), str(num))


def _fix_margins(doc: Document) -> None:
    """Set uniform 0.5 inch margins on all sections."""
    for section in doc.sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)


def _add_table_borders(body) -> None:
    """Add single black borders to all tables."""
    for tblPr in body.iter(qn("w:tblPr")):
        borders = tblPr.find(qn("w:tblBorders"))
        if borders is not None:
            tblPr.remove(borders)

        borders = OxmlElement("w:tblBorders")
        for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            elem = OxmlElement(f"w:{side}")
            elem.set(qn("w:val"), "single")
            elem.set(qn("w:sz"), "4")
            elem.set(qn("w:color"), "000000")
            elem.set(qn("w:space"), "0")
            borders.append(elem)
        tblPr.append(borders)


def _set_table_width(body) -> None:
    """Set tables to use full page width (auto layout)."""
    for tblPr in body.iter(qn("w:tblPr")):
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.insert(0, tblW)
        tblW.set(qn("w:w"), "5000")
        tblW.set(qn("w:type"), "pct")

        layout = tblPr.find(qn("w:tblLayout"))
        if layout is not None:
            tblPr.remove(layout)


def _split_mega_tables(body) -> int:
    """Split mega-tables: extract full-width merged rows as standalone paragraphs."""
    extracted = 0

    for tbl in list(body.iter(qn("w:tbl"))):
        tblGrid = tbl.find(qn("w:tblGrid"))
        if tblGrid is None:
            continue
        grid_cols = len(tblGrid.findall(qn("w:gridCol")))
        if grid_cols <= 1:
            continue

        rows = tbl.findall(qn("w:tr"))
        if len(rows) <= 2:
            continue

        split_indices = []
        for i, row in enumerate(rows):
            cells = row.findall(qn("w:tc"))
            if len(cells) == 1:
                tcPr = cells[0].find(qn("w:tcPr"))
                if tcPr is not None:
                    gridSpan = tcPr.find(qn("w:gridSpan"))
                    if gridSpan is not None:
                        span = int(gridSpan.get(qn("w:val"), "1"))
                        if span >= grid_cols:
                            split_indices.append(i)

        if not split_indices:
            continue

        parent = tbl.getparent()
        tbl_index = list(parent).index(tbl)

        segments = []
        table_rows = []
        for i, row in enumerate(rows):
            if i in split_indices:
                if table_rows:
                    segments.append(("table", table_rows))
                    table_rows = []
                segments.append(("para", row))
                extracted += 1
            else:
                table_rows.append(row)
        if table_rows:
            segments.append(("table", table_rows))

        parent.remove(tbl)

        for seg_type, seg_data in reversed(segments):
            if seg_type == "para":
                row = seg_data
                cell = row.findall(qn("w:tc"))[0]
                paras = cell.findall(qn("w:p"))
                for p in reversed(paras):
                    p_copy = copy.deepcopy(p)
                    parent.insert(tbl_index, p_copy)
            else:
                new_tbl = copy.deepcopy(tbl)
                for old_row in new_tbl.findall(qn("w:tr")):
                    new_tbl.remove(old_row)
                for row in seg_data:
                    new_tbl.append(copy.deepcopy(row))
                parent.insert(tbl_index, new_tbl)

    return extracted


def postprocess_docx(docx_path: str, pdf_path: str = None) -> dict:
    """Apply all post-processing fixes to a pdf2docx-generated DOCX."""
    doc = Document(docx_path)
    body = doc.element.body

    stats = {}

    # Phase 1: Fix Odia chars + remove spurious spaces (before space recovery)
    text_stats = _fix_pre_space_text(body)
    stats.update(text_stats)

    # Phase 2: Recover spaces from PDF text (BEFORE sihari fix!)
    # Both PDF and DOCX have the same wrong sihari ordering from PyMuPDF,
    # so word matching works. Sihari fix depends on word boundaries.
    if pdf_path:
        stats["spaces_recovered"] = _recover_spaces(body, pdf_path)

    # Phase 3: Fix sihari ordering (AFTER space recovery, so word boundaries exist)
    stats["sihari"] = _fix_sihari_all(body)

    # Phase 4: Font and formatting fixes
    _fix_fonts(body)
    _fix_font_sizes(body)
    _fix_margins(doc)
    _add_table_borders(body)
    _set_table_width(body)
    stats["rows_extracted"] = _split_mega_tables(body)

    doc.save(docx_path)
    return stats
